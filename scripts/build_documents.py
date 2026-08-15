from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

NAVY = "17324D"
GREEN = "087A5B"
MINT = "E7F4EF"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
MUTED = "637381"
WHITE = "FFFFFF"


def font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, color):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd")) or OxmlElement("w:shd")
    if node.getparent() is None:
        props.append(node)
    node.set(qn("w:fill"), color)


def margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    node = props.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        props.append(node)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        edge = node.find(qn(f"w:{key}")) or OxmlElement(f"w:{key}")
        if edge.getparent() is None:
            node.append(edge)
        edge.set(qn("w:w"), str(value))
        edge.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    props = table._tbl.tblPr
    width = props.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        props.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = props.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        props.append(indent)
    first_cell_props = table.rows[0].cells[0]._tc.get_or_add_tcPr()
    cell_margins = first_cell_props.find(qn("w:tcMar"))
    start_margin = cell_margins.find(qn("w:start")) if cell_margins is not None else None
    indent.set(qn("w:w"), start_margin.get(qn("w:w")) if start_margin is not None else "120")
    indent.set(qn("w:type"), "dxa")
    row_props = table.rows[0]._tr.get_or_add_trPr()
    header = row_props.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        row_props.append(header)
    header.set(qn("w:val"), "true")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell.width = Inches(value / 1440)
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(value))
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:type"), "dxa")


def configure(doc, title, proposal=False):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.49)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_after = Pt(8 if proposal else 6)
    normal.paragraph_format.line_spacing = 1.333 if proposal else 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if proposal else WD_ALIGN_PARAGRAPH.LEFT
    specs = (("Title", 29, NAVY, 0, 12), ("Heading 1", 16, GREEN, 18, 10),
             ("Heading 2", 13, GREEN, 12 if proposal else 14, 6 if proposal else 7),
             ("Heading 3", 12, NAVY, 8 if proposal else 10, 4 if proposal else 5))
    for name, size, color, before, after in specs:
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.bold, style.font.color.rgb = True, RGBColor.from_string(color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    header = sec.header.paragraphs[0]
    header.text = "SAÚDE PERTO DE VOCÊ  |  ALTAIR/SP"
    font(header.runs[0], 8, MUTED, True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("Material demonstrativo • Agosto de 2026"), 8, MUTED)
    doc.core_properties.title = title
    doc.core_properties.subject = "Transformação digital da saúde pública municipal"
    doc.core_properties.author = "Saúde Perto de Você"


def cover(doc, kicker, title, subtitle, proposal=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82 if proposal else 105)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if proposal else WD_ALIGN_PARAGRAPH.LEFT
    font(p.add_run(kicker.upper()), 11, GREEN, True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if proposal else WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if proposal else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(28)
    font(p.add_run(subtitle), 14, MUTED)
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    shade(cell, NAVY)
    margins(cell, 240, 260, 240, 260)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run("Cuidado simples para quem precisa. Gestão clara para quem atende."), 17, WHITE, True)
    p = cell.add_paragraph("Medicamentos, agendamentos, especialidades e comunicação em uma experiência inclusiva e rastreável.")
    p.paragraph_format.space_after = Pt(0)
    font(p.runs[0], 10, "CDE2DA")
    set_table_geometry(box, [9360])
    p = doc.add_paragraph("Versão demonstrativa • Prefeitura de Altair/SP")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if proposal else WD_ALIGN_PARAGRAPH.LEFT
    font(p.runs[0], 9, MUTED)
    doc.add_page_break()


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, text in enumerate(headers):
        c = t.rows[0].cells[idx]
        shade(c, LIGHT)
        margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), 9, NAVY, True)
    for row in rows:
        cells = t.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cells[idx])
            if len(t.rows) % 2 == 1:
                shade(cells[idx], PALE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(text)), 9, NAVY)
    set_table_geometry(t, widths)
    return t


def bullets(doc, items, numbered=False):
    for item in items:
        doc.add_paragraph(item, style="List Number" if numbered else "List Bullet")


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    shade(c, MINT)
    margins(c, 170, 210, 170, 210)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(label), 10, GREEN, True)
    p = c.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    font(p.runs[0], 10, NAVY)
    set_table_geometry(t, [9360])


def manual():
    doc = Document()
    configure(doc, "Manual de uso — Saúde Perto de Você")
    cover(doc, "Guia de uso", "Saúde Perto de Você", "Manual da demonstração para cidadãos e equipes municipais")
    doc.add_heading("1. Comece por aqui", 1)
    doc.add_paragraph("O Saúde Perto de Você reúne o estoque público de medicamentos, solicitações, retirada ou entrega agendada, demanda por especialistas e comunicação com a farmácia. A navegação do cidadão foi desenhada para celular, com linguagem direta e poucos passos; o painel administrativo prioriza informação operacional e indicadores.")
    doc.add_heading("Endereço e acessos", 2)
    doc.add_paragraph("Abra o endereço público informado na entrega. Todos os dados abaixo são fictícios e exclusivos da demonstração.")
    table(doc, ["Cenário", "Identificação", "Senha / código"], [
        ("Cidadã validada", "CPF 123.456.789-09", "Cidadao@2026"),
        ("Nova ativação", "CPF 111.222.333-44 • 18/06/1972", "SAUDE-2026"),
        ("Admin da farmácia", "admin@altair.sp.gov.br", "Admin@2026"),
    ], [1900, 4000, 3460])
    callout(doc, "Cenário carregado", "24 medicamentos, 15 cidadãos, lotes e validades, solicitações em múltiplas etapas, quatro especialidades, consultas previstas, mensagens e avisos.")

    doc.add_heading("2. Portal público", 1)
    bullets(doc, [
        "Pesquise o medicamento pelo nome, princípio ativo ou código.",
        "Confira disponibilidade, apresentação, necessidade de receita e data de atualização.",
        "Acesse o portal do cidadão ou o painel da farmácia pelo botão de entrada.",
    ], numbered=True)
    doc.add_paragraph("A consulta pública apoia o atendimento à obrigação de transparência prevista na Lei nº 14.654/2023. A data exibida vem da última movimentação real do inventário, e não do momento da consulta. O município continua responsável por alimentar, conferir e publicar o estoque pelo menos a cada quinze dias.")

    doc.add_heading("3. Ativação e primeiro acesso", 1)
    bullets(doc, [
        "Selecione Primeiro acesso e informe CPF, data de nascimento e código entregue pela prefeitura.",
        "Crie uma senha com pelo menos dez caracteres, combinando letras e números.",
        "Envie os documentos solicitados pelo portal.",
        "Na primeira retirada, compareça presencialmente para conferência documental.",
        "A entrega domiciliar só é liberada depois de duas etapas: documentos validados e primeira retirada presencial concluída.",
    ], numbered=True)
    callout(doc, "Teste recomendado", "Use Ana Lúcia Souza para percorrer a ativação: CPF 111.222.333-44, nascimento 18/06/1972 e código SAUDE-2026.")

    doc.add_heading("4. Jornada do cidadão", 1)
    doc.add_heading("Solicitar medicamento", 2)
    bullets(doc, [
        "Entre como Maria Aparecida e escolha Solicitar medicamento.",
        "Selecione o produto e a quantidade; anexe a receita quando necessário.",
        "Escolha retirada ou entrega. Retiradas usam intervalos de 15 minutos; entregas usam faixas de uma hora e exigem ao menos 24 horas de antecedência.",
        "Revise e envie. O protocolo aparece imediatamente para acompanhamento.",
    ], numbered=True)
    doc.add_heading("Quando não houver saldo", 2)
    doc.add_paragraph("O cidadão registra a necessidade em um toque. Quando um novo lote for importado ou lançado e o saldo voltar a ficar disponível, o sistema cria um aviso no portal.")
    doc.add_heading("Consultas e especialidades", 2)
    doc.add_paragraph("Em Consultar especialidades, o cidadão vê os próximos atendimentos e registra uma única intenção por especialidade. Quando a prefeitura cria uma agenda, os interessados recebem os detalhes de data, horário e local e podem confirmar presença enquanto houver capacidade.")
    doc.add_heading("Tela inicial", 2)
    doc.add_paragraph("A página principal destaca consultas pendentes de confirmação e já confirmadas, pedidos recentes, notificações e atalhos. O cidadão pode ativar Web Push em cada aparelho para receber alertas de estoque, solicitações, consultas e mensagens. O chat mantém a conversa com a farmácia dentro do sistema, sem depender de WhatsApp.")

    doc.add_heading("5. Painel da farmácia", 1)
    table(doc, ["Área", "Uso principal"], [
        ("Visão geral", "KPIs de pedidos, estoque baixo, atendimentos e cidadãos."),
        ("Solicitações", "Analisar, aprovar, preparar, entregar/retirar e concluir."),
        ("Estoque", "Produtos, lotes, validade, saldo, reserva e movimentação."),
        ("Agenda", "Horários, capacidade simultânea, retirada e entrega."),
        ("Especialidades", "Demanda, escalas, confirmação e situação da consulta."),
        ("Cidadãos", "Cadastro, validação, documentos e importação em massa."),
        ("Chat", "Mensagens e avisos vinculados ao cidadão."),
        ("Privacidade", "Pedidos do titular: acesso, correção, informação e revisão."),
        ("Auditoria", "Eventos sensíveis visíveis a gestores autorizados."),
        ("Equipe", "Usuários por perfil: operador, gestor e administrador geral."),
    ], [2050, 7310])
    doc.add_heading("Fluxo de uma solicitação", 2)
    bullets(doc, [
        "Recebida: conferir cidadão, receita, produto, quantidade e modalidade.",
        "Aprovada: o sistema reserva um lote específico pelo critério FEFO, priorizando a validade mais próxima e não vencida.",
        "Pronta: informar ao cidadão que a retirada ou entrega foi preparada.",
        "Concluída: consumir exatamente o lote reservado, registrar a saída física e encerrar o atendimento.",
        "Não aprovada: exigir motivo, avisar o cidadão e liberar horário e reserva de estoque quando aplicável.",
    ], numbered=True)

    doc.add_heading("6. Agendas e capacidade", 1)
    doc.add_paragraph("A equipe cadastra previamente os horários disponíveis. A capacidade máxima por faixa é configurável e o painel mostra quantos agendamentos já ocupam cada período, reduzindo concentração e filas.")
    table(doc, ["Modalidade", "Granularidade", "Regra do MVP"], [
        ("Retirada", "15 minutos", "Capacidade simultânea configurável."),
        ("Entrega", "1 hora", "Mínimo de 24 horas de antecedência."),
        ("Especialista", "Data e hora da escala", "Local obrigatoriamente cadastrado."),
    ], [1800, 1900, 5660])
    doc.add_paragraph("A situação da consulta segue transições controladas: planejada pode ser confirmada ou cancelada; confirmada pode ser concluída ou cancelada; concluída e cancelada são estados finais. Cada mudança gera aviso aos cidadãos vinculados.")

    doc.add_heading("7. Importações e cadastros", 1)
    table(doc, ["Carga", "Campos mínimos"], [
        ("Inventário CSV", "codigo, nome, lote, validade, saldo"),
        ("Cidadãos CSV", "cpf, nome, nascimento, endereco, bairro, codigo_ativacao"),
        ("XML de produto", "cProd, xProd, qCom e, quando disponíveis, nLote, qLote, dVal"),
    ], [1900, 7460])
    doc.add_paragraph("Produtos inexistentes são criados durante a importação do XML; a equipe também pode cadastrá-los manualmente. Toda ação sensível deixa trilha de auditoria.")

    doc.add_heading("8. Roteiro de demonstração", 1)
    bullets(doc, [
        "Abra a consulta pública e pesquise Losartana e Azitromicina para contrastar saldo disponível e item zerado.",
        "Ative o acesso de Ana Lúcia e mostre a regra da primeira retirada presencial.",
        "Entre como Maria, solicite um medicamento e escolha retirada ou entrega.",
        "Mostre os cards de consultas pendentes e confirmadas e confirme uma agenda.",
        "Entre como admin, percorra KPIs, solicitações, lotes, horários e demanda por especialidade.",
        "Altere a situação de uma consulta e mostre a notificação correspondente.",
        "Encerre no portal público e destaque a Lei nº 14.654/2023.",
    ], numbered=True)
    doc.add_heading("9. Privacidade, segurança e compliance", 1)
    bullets(doc, [
        "Dados de saúde são sensíveis: o acesso é autenticado, o CPF aparece mascarado e senhas/códigos são protegidos por hash.",
        "Uploads aceitam apenas PDF, JPG ou PNG de até 5 MB e validam o conteúdo real do arquivo.",
        "Cinco falhas de login bloqueiam novas tentativas por quinze minutos.",
        "Operadores cuidam da rotina; gestores administram estoque, agendas e privacidade; somente o administrador geral cria acessos.",
        "O cidadão pode registrar pedidos de acesso, correção, informação ou revisão na área Privacidade.",
        "A consulta pública nunca exibe cidadãos, pedidos, receitas ou dados clínicos.",
    ])
    callout(doc, "Atenção jurídica", "O sistema fornece controles técnicos, mas não certifica conformidade isoladamente. Antes da produção, o município deve definir controlador, encarregado, bases legais por finalidade, tabela de retenção, RIPD quando aplicável, contratos, resposta a incidentes e canal alternativo acessível.")
    doc.add_heading("10. Limites e pré-requisitos de produção", 1)
    doc.add_paragraph("A versão demonstra experiência, regras e integração dos fluxos com dados fictícios. Os avisos ficam registrados no portal e Web Push funciona quando as chaves VAPID estão configuradas e o cidadão autoriza o aparelho. Para produção municipal são necessários banco e armazenamento persistentes, domínio oficial, identidade visual, rotina de atualização quinzenal comprovável, backup e restauração testados, monitoramento, homologação de segurança, avaliação eMAG/WCAG com pessoas com deficiência, plano de incidentes, política LGPD e treinamento da equipe.")
    path = OUT / "Manual-Saude-Perto-de-Voce.docx"
    doc.save(path)
    return path


def pitch():
    doc = Document()
    configure(doc, "Pitch comercial — Saúde Perto de Você para Altair/SP", proposal=True)
    cover(doc, "Proposta municipal", "Saúde Perto de Você", "Medicamentos e consultas mais perto do cidadão — gestão orientada por demanda", proposal=True)
    doc.add_heading("A ideia central", 1)
    p = doc.add_paragraph()
    font(p.add_run("A prefeitura deixa de organizar a saúde apenas pela fila e passa a organizar pelo que a população realmente precisa."), 15, NAVY, True)
    doc.add_paragraph("O cidadão consulta estoque, agenda retirada ou entrega, registra necessidade de especialista e recebe informações claras. A equipe ganha previsibilidade, rastreabilidade e indicadores para planejar atendimento, compras e escalas.")

    doc.add_heading("O problema que resolvemos", 1)
    bullets(doc, [
        "Deslocamentos sem garantia de que o medicamento está disponível.",
        "Filas concentradas e pouca previsibilidade de retirada.",
        "Estoque, lotes e validades dispersos em controles manuais.",
        "Demanda por especialistas conhecida tarde demais ou apenas por relatos.",
        "Comunicação fragmentada entre cidadão e farmácia.",
        "Dificuldade para cumprir e comprovar a transparência do estoque público.",
    ])
    callout(doc, "Obrigação transformada em serviço", "A Lei nº 14.654/2023 exige publicação eletrônica e acessível dos estoques das farmácias públicas do SUS, com atualização ao menos quinzenal. A proposta apoia essa obrigação com data da última movimentação real e uma visão operacional mais frequente.")

    doc.add_heading("A solução", 1)
    table(doc, ["Frente", "Entrega"], [
        ("Transparência", "Portal público de estoque, pesquisa simples e data da atualização."),
        ("Acesso", "PWA mobile-first, ativação segura e primeira validação presencial."),
        ("Medicamentos", "Solicitação, retirada ou entrega, agenda, lotes, validade e FEFO."),
        ("Especialidades", "Intenção de consulta, mapa de demanda, escalas e confirmações."),
        ("Gestão", "KPIs, perfis segregados, importações, auditoria e fila LGPD."),
        ("Comunicação", "Chat, avisos no portal e Web Push opcional por aparelho, sem WhatsApp ou SMS no escopo inicial."),
    ], [1900, 7460])

    doc.add_heading("Valor para cada público", 1)
    table(doc, ["Público", "Valor percebido"], [
        ("Cidadão", "Menos viagens perdidas, menos fila e informação simples no celular."),
        ("Farmácia", "Fila organizada, lote correto, comunicação e trabalho rastreável."),
        ("Secretaria", "Demanda real por medicamento, bairro, modalidade e especialidade."),
        ("Gestão municipal", "Decisões com dados, transparência legal e visão de capacidade."),
        ("Controle interno", "Histórico de acessos, reservas, saídas e alterações."),
    ], [2000, 7360])

    doc.add_heading("Por que começar agora", 1)
    doc.add_paragraph("O MVP concentra os fluxos de maior impacto sem exigir aplicativo em lojas, WhatsApp ou integrações complexas. A PWA reduz atrito de adoção e permite validar rapidamente linguagem, capacidade operacional, logística de entrega e regras de atendimento.")
    doc.add_heading("Escopo recomendado do piloto", 2)
    bullets(doc, [
        "Uma farmácia municipal e locais de saúde previamente cadastrados.",
        "Carga inicial de cidadãos, produtos, lotes, validades e saldos.",
        "Retirada em blocos de 15 minutos e entrega em faixas de uma hora.",
        "Quatro especialidades prioritárias com mapa de demanda.",
        "Treinamento da equipe, acompanhamento de adoção e reunião semanal de ajustes.",
    ])

    doc.add_heading("Indicadores do piloto", 1)
    table(doc, ["Indicador", "O que demonstra"], [
        ("Consultas públicas ao estoque", "Informação chegando antes do deslocamento."),
        ("Tempo e volume por faixa", "Redução de fila e equilíbrio da capacidade."),
        ("Pedidos atendidos", "Efetividade da farmácia e disponibilidade real."),
        ("Alertas de reposição", "Demanda reprimida e resposta ao reabastecimento."),
        ("Intenções por especialidade", "Base objetiva para concentrar escalas médicas."),
        ("Confirmação de consultas", "Aderência à agenda criada pelo município."),
        ("Atualização do inventário", "Evidência do cumprimento do ciclo quinzenal."),
        ("Solicitações de privacidade", "Capacidade de atender e rastrear direitos do titular."),
    ], [3400, 5960])

    doc.add_heading("Evolução independente: teleconsulta", 1)
    doc.add_paragraph("Como proposta complementar, a teleconsulta pode ampliar cobertura com custo menor, especialmente enquanto a demanda local não justifica a presença frequente de determinado especialista. Ela deve ser contratada e avaliada separadamente do MVP, preservando foco e clareza de implantação.")

    doc.add_heading("Próximo passo", 1)
    doc.add_paragraph("Realizar uma reunião de validação com Saúde, Farmácia, Atenção Básica, TI e Controle Interno. Em seguida, fechar dados iniciais, regras de capacidade, bairros atendidos por entrega, documentos de validação, responsáveis e indicadores dos primeiros 60 dias.")
    callout(doc, "Proposta de decisão", "Aprovar um piloto assistido, medir adesão e impacto operacional e então expandir gradualmente para toda a rede municipal.")
    doc.add_heading("Governança necessária", 2)
    doc.add_paragraph("A contratação deve incluir papéis de controlador e operador, requisitos de segurança, retenção, backup, continuidade, resposta a incidentes, acessibilidade e homologação. O produto entrega prontidão técnica; a conformidade final depende também de atos, processos e responsáveis do município.")
    path = OUT / "Pitch-Comercial-Saude-Perto-de-Voce-Altair.docx"
    doc.save(path)
    return path


def compliance_report():
    doc = Document()
    configure(doc, "Revisão de regras de negócio e compliance — Saúde Perto de Você")
    cover(doc, "Relatório de revisão", "Regras de negócio e compliance", "Controles técnicos, lacunas organizacionais e critérios para implantação municipal")
    doc.add_heading("1. Parecer executivo", 1)
    callout(doc, "Conclusão", "O MVP está tecnicamente mais seguro e coerente após a revisão, mas não deve ser apresentado como produto juridicamente certificado. A conformidade final exige controles técnicos e também decisões formais do Município de Altair/SP.")
    doc.add_paragraph("Foram revisados portal público, autenticação, jornada do cidadão, solicitação e dispensação de medicamentos, estoque/lotes, agendas, especialistas, usuários administrativos, documentos, privacidade, auditoria, dependências e mensagens de interface. Os dados permanecem fictícios.")
    table(doc, ["Dimensão", "Situação após correções", "Condição para produção"], [
        ("Transparência de estoque", "Pronto tecnicamente", "Rotina municipal que garanta conferência e atualização em até 15 dias."),
        ("Regras de medicamentos", "Fluxo corrigido", "Validação farmacêutica, protocolos locais e responsáveis definidos."),
        ("LGPD e dados de saúde", "Controles técnicos parciais", "Encarregado, bases legais, inventário, retenção, RIPD e contratos."),
        ("Segurança", "Controles essenciais aplicados", "Pentest, monitoramento, backup, continuidade e resposta a incidentes."),
        ("Acessibilidade", "Base semântica e mobile", "Auditoria eMAG/WCAG, testes por teclado/leitor e com usuários."),
        ("Demonstração pública", "Segura para dados fictícios", "Não inserir dados reais no ambiente demonstrativo."),
    ], [2100, 3000, 4260])

    doc.add_heading("2. Correções implementadas", 1)
    table(doc, ["Risco encontrado", "Correção aplicada", "Resultado esperado"], [
        ("Reserva duplicada", "Transição recebida → aprovada validada e reserva atômica.", "Repetir a aprovação não duplica saldo reservado."),
        ("Lote errado na conclusão", "Pedido guarda reserved_lot_id e consome exatamente esse lote.", "Rastreabilidade entre cidadão, protocolo, lote e saída."),
        ("Estados pulados", "Máquina de estados para pedido e consulta.", "Não é possível concluir antes de aprovar e preparar."),
        ("Recusa sem justificativa", "Motivo obrigatório, aviso ao cidadão e liberação de recursos.", "Decisão explicável e fila/estoque consistentes."),
        ("Concorrência no horário", "Trigger de banco valida capacidade e incrementa reserva no INSERT.", "Duas solicitações simultâneas não ultrapassam a capacidade."),
        ("Entrega liberada cedo", "Validação documental e primeira retirada são etapas independentes.", "Entrega somente após ambas concluídas."),
        ("Upload permissivo", "Allowlist PDF/JPG/PNG, limite de 5 MB e assinatura mágica.", "Redução de arquivo disfarçado e carga excessiva."),
        ("Tentativas ilimitadas", "Bloqueio por identidade após cinco falhas durante 15 minutos.", "Mitigação de força bruta."),
        ("Perfis sem separação", "RBAC: operador, gestor e administrador geral.", "Privilégio mínimo por função."),
        ("Sem canal LGPD", "Pedidos de acesso, correção, informação e revisão.", "Demanda do titular registrada e acompanhável."),
        ("Data de estoque artificial", "Portal usa a última atualização real dos lotes.", "Transparência verificável."),
        ("Dependências vulneráveis", "Next atualizado e auditoria de produção zerada.", "0 vulnerabilidades conhecidas no npm audit --omit=dev em 15/08/2026."),
    ], [2200, 3900, 3260])

    doc.add_heading("3. Regras de negócio consolidadas", 1)
    doc.add_heading("Cadastro, ativação e entrega", 2)
    bullets(doc, [
        "Ativação: CPF, data de nascimento, código vigente e senha com dez ou mais caracteres, contendo letras e números.",
        "CPF é usado normalizado e armazenado como hash; somente a versão mascarada aparece nas telas administrativas.",
        "Código de ativação é de uso único e possui validade; um novo código revoga o anterior.",
        "Validar documentos não elimina a obrigação da primeira retirada presencial.",
        "Entrega domiciliar exige cadastro validado, primeira retirada concluída, produto elegível e faixa com pelo menos 24 horas de antecedência.",
    ])
    doc.add_heading("Solicitação e dispensação", 2)
    table(doc, ["Estado atual", "Próximos estados permitidos", "Efeito"], [
        ("Recebida", "Aprovada ou não aprovada", "Ocupa horário; aguarda análise."),
        ("Aprovada", "Pronta/programada ou não aprovada", "Reserva lote FEFO específico."),
        ("Pronta/programada", "Concluída ou não aprovada", "Aguarda retirada/entrega."),
        ("Concluída", "Nenhum", "Baixa física e reservada no lote vinculado."),
        ("Não aprovada", "Nenhum", "Registra motivo e libera horário/reserva futura."),
    ], [2100, 3000, 4260])
    bullets(doc, [
        "Receita é obrigatória quando o cadastro do produto assim definir.",
        "O saldo público não constitui promessa de dispensação; cadastro, receita, quantidade e lote são conferidos.",
        "Lotes vencidos não podem ser escolhidos na aprovação.",
        "Sem saldo, o cidadão registra uma necessidade única por medicamento e recebe aviso quando uma nova entrada torna o item disponível.",
    ])
    doc.add_heading("Agenda de especialistas", 2)
    bullets(doc, [
        "A manifestação é uma intenção para planejamento; não é encaminhamento, regulação ou consulta confirmada.",
        "Uma intenção ativa por cidadão e especialidade evita contagem duplicada.",
        "A agenda usa somente local cadastrado e ativo, data futura e capacidade entre 1 e 500.",
        "Confirmação do cidadão é aceita somente em agenda confirmada, futura e com vaga.",
        "Consulta planejada pode ser confirmada ou cancelada; confirmada pode ser concluída ou cancelada; estados finais não reabrem automaticamente.",
    ])

    doc.add_heading("4. Matriz de conformidade", 1)
    table(doc, ["Referência", "Exigência aplicada ao produto", "Avaliação"], [
        ("Lei 14.654/2023", "Estoque das farmácias públicas em página eletrônica, acessível ao cidadão comum e atualizado ao menos quinzenalmente.", "Atendido tecnicamente; governança da atualização é municipal."),
        ("LGPD, arts. 5º, 7º, 11 e 23", "Saúde é dado sensível; tratamento público requer finalidade, base legal, transparência e necessidade.", "Parcial: minimização e canal implementados; atos e inventário pendentes."),
        ("LGPD, arts. 37, 41, 46, 48 e 50", "Registros, encarregado, segurança, incidentes e governança.", "Parcial: auditoria e segurança aplicadas; programa formal pendente."),
        ("Resolução CD/ANPD 15/2024", "Processo para avaliar e comunicar incidentes de segurança.", "Pendente organizacional: plano, responsáveis e simulado."),
        ("Lei 13.146/2015, art. 63", "Acessibilidade obrigatória em sítios de órgãos de governo.", "Parcial: semântica e responsividade; homologação assistiva pendente."),
        ("eMAG 3.1 / WCAG", "Padrões Web, teclado, labels, instruções, contraste e avaliação.", "Parcial: vários fundamentos presentes; auditoria completa pendente."),
    ], [2100, 4300, 2960])
    callout(doc, "Sem certificação automática", "Nenhum teste de software substitui parecer jurídico, farmacêutico, de segurança ou acessibilidade. Este relatório classifica prontidão técnica e tarefas de implantação.")

    doc.add_heading("5. LGPD: mapa de dados e finalidades", 1)
    table(doc, ["Dados", "Finalidade", "Proteção/observação"], [
        ("CPF e nascimento", "Identificar e ativar o cidadão.", "CPF com hash e máscara; restringir consulta bruta."),
        ("Endereço/bairro", "Cadastro e entrega domiciliar.", "Usar somente quando necessário; definir retenção."),
        ("Receitas/documentos", "Validar direito à dispensação e identidade.", "Armazenamento privado, allowlist e acesso restrito."),
        ("Medicamentos e pedidos", "Atender política pública de assistência farmacêutica.", "Dado de saúde sensível; auditoria e mínimo acesso."),
        ("Interesse em especialista", "Planejar oferta municipal.", "Não expor demanda individual no portal público."),
        ("Mensagens e avisos", "Comunicação operacional.", "Orientar equipe a evitar dados desnecessários."),
        ("Logs de auditoria", "Segurança, responsabilização e controle.", "Definir retenção, acesso e proteção contra alteração."),
    ], [1900, 3500, 3960])

    doc.add_heading("6. Ações obrigatórias antes do go-live", 1)
    bullets(doc, [
        "Nomear/publicar controlador e contato do encarregado; definir canal alternativo presencial/telefônico.",
        "Registrar operações, finalidades, categorias, bases legais, compartilhamentos e prazos de retenção.",
        "Elaborar RIPD para o tratamento de dados sensíveis e riscos de entrega/integração, quando indicado.",
        "Formalizar contrato de operador/suboperadores, hospedagem, localização dos dados, descarte e auditoria.",
        "Configurar domínio oficial, TLS, segredos, banco/R2 de produção e ambientes separados de teste.",
        "Testar backup/restauração, continuidade, monitoramento, alertas e resposta a incidentes conforme RCIS.",
        "Executar pentest independente e corrigir achados críticos/altos antes de dados reais.",
        "Realizar auditoria eMAG/WCAG, navegação por teclado, leitor de tela, contraste, zoom e teste com cidadãos idosos e com deficiência.",
        "Definir protocolo farmacêutico: receita, quantidade, substituição, controlados, cadeia fria, entrega e comprovação.",
        "Treinar operadores e gestores em LGPD, engenharia social, arquivos, atendimento e contingência.",
    ], numbered=True)

    doc.add_heading("7. Critérios de aceite do piloto", 1)
    table(doc, ["Teste", "Critério de aprovação"], [
        ("Estoque público", "Data real, escopo da unidade e atualização monitorada em até 15 dias."),
        ("Concorrência", "Capacidade nunca excedida em solicitações simultâneas."),
        ("Lote", "Reserva e baixa usam o mesmo lote; saldo físico/reservado permanece não negativo."),
        ("Privacidade", "Operador não acessa auditoria/fila LGPD; cidadão vê somente os próprios dados."),
        ("Arquivos", "Formato disfarçado, arquivo >5 MB e extensão não permitida são bloqueados."),
        ("Acessibilidade", "Fluxos críticos completos por teclado e leitor de tela, sem perda a 200% de zoom."),
        ("Incidente", "Equipe executa simulado, classifica risco e aciona responsáveis dentro do procedimento."),
        ("Continuidade", "Restauração comprovada com RPO/RTO aprovados pelo município."),
    ], [2400, 6960])

    doc.add_heading("8. Referências oficiais", 1)
    bullets(doc, [
        "Lei nº 14.654/2023: https://www.planalto.gov.br/ccivil_03/_ato2023-2026/2023/lei/l14654.htm",
        "Lei Geral de Proteção de Dados (Lei nº 13.709/2018): https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709compilado.htm",
        "Guia ANPD - Tratamento de dados pessoais pelo Poder Público: https://www.gov.br/anpd/pt-br/centrais-de-conteudo/materiais-educativos-e-publicacoes/guia_tratamento_de_dados_pessoais_pelo_poder_publico___defeso_eleitoral.pdf",
        "Resolução CD/ANPD nº 15/2024: https://www.gov.br/anpd/pt-br/assuntos/noticias/anpd-aprova-o-regulamento-de-comunicacao-de-incidente-de-seguranca",
        "Lei Brasileira de Inclusão (Lei nº 13.146/2015): https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2015/lei/l13146.htm",
        "Modelo de Acessibilidade em Governo Eletrônico (eMAG): https://www.gov.br/governodigital/pt-br/acessibilidade-e-usuario/acessibilidade-digital/modelo-de-acessibilidade",
    ])
    path = OUT / "Revisao-Regras-Negocio-Compliance-Saude-Perto-de-Voce.docx"
    doc.save(path)
    return path


def presentation_script():
    doc = Document()
    configure(doc, "Roteiro detalhado do produto — Saúde Perto de Você")
    cover(doc, "Roteiro para apresentação", "Saúde Perto de Você", "Narrativa cena a cena para criação do deck comercial e demonstração do produto")
    doc.add_heading("Como usar este roteiro", 1)
    doc.add_paragraph("Cada bloco corresponde a um slide ou cena. A apresentação principal foi planejada para 20 a 25 minutos, seguida por 10 minutos de demonstração e perguntas. O roteiro separa mensagem, visual, fala e prova para facilitar a criação posterior do deck.")
    table(doc, ["Formato", "Duração", "Uso"], [
        ("Executivo", "8-10 min", "Cenas 1 a 5, 13, 15 e 16."),
        ("Comercial", "20-25 min", "Todas as cenas, com demonstração resumida."),
        ("Técnico/compliance", "30-40 min", "Todas as cenas, aprofundando 11, 12 e 14."),
    ], [1900, 1800, 5660])

    scenes = [
        ("1. Abertura - saúde sem viagem perdida", "Criar identificação imediata com o problema cotidiano.", "Tela limpa com celular exibindo a consulta de estoque; ao fundo, mapa simples do município.", "Antes de sair de casa, o cidadão deveria saber se o medicamento está disponível, quando pode retirar e onde será atendido.", "Começar pela experiência humana, não pela tecnologia.", "E se o município pudesse transformar essa incerteza em uma jornada simples?"),
        ("2. O problema municipal", "Mostrar que filas e falta de dados têm a mesma raiz.", "Quatro sinais: deslocamento sem garantia, fila concentrada, estoque disperso, demanda de especialistas invisível.", "Hoje a equipe reage a filas, planilhas e relatos. O cidadão descobre o problema somente quando chega à unidade.", "Use exemplos locais sem expor pessoas reais.", "O produto organiza essa realidade em três superfícies conectadas."),
        ("3. Visão do produto", "Apresentar a arquitetura funcional.", "Fluxo: Portal público → Portal do cidadão → Painel da farmácia; especialidades e notificações atravessam os três.", "Saúde Perto de Você é uma PWA municipal: funciona no navegador e foi desenhada primeiro para o celular do cidadão e para o computador da equipe.", "Destacar que não há instalação em loja nem WhatsApp no escopo inicial.", "A primeira superfície resolve uma obrigação e um problema ao mesmo tempo."),
        ("4. Transparência de estoque", "Demonstrar valor público e base legal.", "Busca por Losartana disponível e Azitromicina sem estoque; data real da última atualização.", "A Lei 14.654/2023 obriga a divulgação eletrônica, acessível e ao menos quinzenal. O portal apoia o cumprimento com atualização operacional e linguagem comum.", "Deixar claro: saldo não garante dispensação e não exibe dados pessoais.", "Quando o cidadão entra, a transparência vira serviço."),
        ("5. Acesso simples e seguro", "Explicar ativação sem criar barreiras.", "CPF + nascimento + código municipal → criação de senha → portal.", "O cidadão pré-cadastrado ativa o acesso com dados conhecidos e um código de uso único. O CPF fica protegido e aparece mascarado.", "Mostrar bloqueio por tentativas e senha de 10 caracteres.", "A confiança continua na primeira retirada presencial."),
        ("6. Primeira retirada e entrega", "Explicar a regra em duas etapas.", "Checklist: documentos validados + primeira retirada concluída = entrega liberada.", "Enviar documento adianta a conferência, mas não substitui a presença inicial. A entrega só é habilitada quando as duas condições forem concluídas.", "Evitar prometer entrega universal; depende de produto e regra municipal.", "Com o cadastro validado, a solicitação acontece em poucos passos."),
        ("7. Solicitar medicamento", "Mostrar UX mobile com o mínimo de cliques.", "Produto → quantidade → retirada/entrega → horário → receita → confirmação.", "Retirada usa blocos de 15 minutos; entrega usa faixas de uma hora e exige 24 horas de antecedência. Medicamentos que exigem receita bloqueiam o envio sem PDF, JPG ou PNG válido.", "Mostrar aviso: solicitação depende de análise.", "A farmácia recebe um processo rastreável, não apenas uma mensagem."),
        ("8. Operação farmacêutica", "Provar consistência de estoque e lote.", "Linha de estados: recebida → aprovada → pronta/programada → concluída; ramificação para não aprovada.", "Na aprovação, o sistema escolhe um lote não vencido pelo FEFO e grava esse lote no pedido. Na conclusão, baixa exatamente o lote reservado.", "Recusa exige motivo e libera reserva/horário.", "A mesma disciplina organiza a capacidade de atendimento."),
        ("9. Agenda sem filas", "Demonstrar gestão de capacidade.", "Calendário com blocos de 15 min, faixas de 1 h, ocupação e limite simultâneo.", "A equipe configura quantas pessoas cabem em cada faixa. Uma proteção no banco impede ultrapassar a capacidade mesmo com solicitações simultâneas.", "Mostrar que capacidade não pode ficar abaixo dos já agendados.", "O painel transforma a agenda em indicadores operacionais."),
        ("10. Dashboard da farmácia", "Mostrar riqueza de informação sem sobrecarregar.", "KPIs: fila, estoque crítico, agendados hoje, cidadãos; tabelas de solicitações e lotes.", "O operador enxerga o que precisa fazer agora. O gestor acompanha capacidade, estoque, importações e tendências.", "Explicar perfis: operador, gestor e administrador geral.", "Além de medicamentos, o município passa a enxergar demanda reprimida."),
        ("11. Especialistas orientados por demanda", "Posicionar intenção como dado de planejamento.", "Ranking por especialidade e bairro; criação de agenda em local cadastrado; aviso aos interessados.", "O cidadão registra apenas uma intenção por especialidade. Isso não é encaminhamento nem consulta. Quando a agenda é criada, recebe data, horário e local e confirma enquanto houver vaga.", "Evitar linguagem clínica ou promessa de atendimento.", "O resultado é concentrar especialistas nas datas de maior necessidade."),
        ("12. Comunicação e avisos", "Mostrar acompanhamento sem ferramenta externa.", "Cards de consulta pendente/confirmada, aviso de reposição, chat interno e ativação do Web Push.", "O cidadão acompanha pedidos, consultas e mensagens no portal e pode ativar alertas em segundo plano no próprio aparelho. A tela bloqueada recebe texto genérico e os detalhes ficam atrás da autenticação.", "Não prometer WhatsApp ou SMS no MVP. O Web Push depende de permissão do cidadão, navegador compatível e chaves VAPID da implantação.", "Ao comunicar, também precisamos proteger os dados."),
        ("13. Privacidade e segurança", "Demonstrar responsabilidade com dados sensíveis.", "Escudo com: hash, cookie seguro, RBAC, upload validado, rate limit, auditoria, fila LGPD.", "Dados de saúde são sensíveis. O produto minimiza exposição, separa perfis, registra eventos e oferece pedidos de acesso, correção, informação e revisão.", "Dizer explicitamente que conformidade final depende da governança municipal.", "Isso nos leva ao plano seguro de implantação."),
        ("14. Compliance e pré-requisitos", "Separar prontidão técnica de obrigações organizacionais.", "Matriz verde/amarela: estoque e fluxo técnico prontos; LGPD, acessibilidade, incidentes e continuidade dependem de implantação.", "Antes de dados reais: encarregado, bases legais, retenção, RIPD quando aplicável, contratos, pentest, backup, plano de incidentes e avaliação eMAG/WCAG.", "Nunca usar selo '100% LGPD'.", "Com essas condições, o município pode iniciar um piloto controlado."),
        ("15. Piloto assistido", "Converter visão em execução.", "Linha de 60 dias: preparação, carga e treinamento, operação assistida, avaliação.", "Começamos por uma farmácia e locais cadastrados, com dados validados, responsáveis nomeados e indicadores semanais.", "Indicadores: atualização, fila, pedidos, ocupação, alertas, intenções, confirmações e privacidade.", "O piloto produz uma decisão objetiva de expansão."),
        ("16. Teleconsulta como proposta independente", "Apresentar evolução sem diluir o MVP.", "Bloco separado: cobertura especializada com custo menor, contratação própria.", "Teleconsulta pode reduzir desassistência quando a demanda não justifica presença frequente, mas deve ser avaliada e contratada separadamente.", "Não misturar preço ou responsabilidade clínica ao escopo atual.", "Encerrar com a decisão solicitada."),
        ("17. Fechamento e chamada para ação", "Pedir um próximo passo concreto.", "Frase final + três decisões: aprovar piloto, nomear equipe, validar dados/regras.", "Cuidado simples para quem precisa, gestão clara para quem atende. A proposta é validar o piloto com Saúde, Farmácia, Atenção Básica, TI e Controle Interno.", "Deixar na tela URL da demonstração e contato comercial.", "Abrir a demonstração ao vivo e as perguntas."),
    ]
    for title, objective, visual, narration, proof, transition in scenes:
        doc.add_heading(title, 1)
        table(doc, ["Elemento", "Detalhamento"], [
            ("Objetivo", objective),
            ("Visual sugerido", visual),
            ("Fala do apresentador", narration),
            ("Prova/cuidado", proof),
            ("Transição", transition),
        ], [1900, 7460])

    doc.add_heading("Roteiro da demonstração ao vivo", 1)
    bullets(doc, [
        "Portal público: buscar Losartana e Azitromicina; apontar escopo e data real da atualização.",
        "Cidadã Maria: mostrar cards de consultas, fazer pedido e destacar receita/aviso de análise.",
        "Especialidades: diferenciar intenção, aviso, confirmação e consulta confirmada.",
        "Privacidade: abrir canal do titular e explicar por que exclusão pode depender de obrigação legal.",
        "Admin Gabriela: mostrar KPIs, fila, lote reservado, horários, agendas, privacidade e auditoria.",
        "Simulação: recusar com motivo ou avançar um pedido; nunca alterar dados reais na apresentação.",
    ], numbered=True)
    doc.add_heading("Perguntas difíceis e respostas", 1)
    table(doc, ["Pergunta", "Resposta recomendada"], [
        ("O estoque garante o remédio?", "Não. É transparência operacional; a dispensação depende de cadastro, receita, quantidade e conferência."),
        ("Está 100% LGPD?", "O produto possui controles técnicos. A conformidade completa exige governança, documentos, responsáveis e operação municipal."),
        ("A intenção garante especialista?", "Não. Ela mede demanda; o agendamento ocorre somente quando uma agenda confirmada é vinculada e há vaga."),
        ("Tem WhatsApp ou push?", "O MVP inclui Web Push opcional e avisos internos. WhatsApp e SMS permanecem fora do escopo inicial."),
        ("Pode usar dados reais nesta demo?", "Não. A demonstração pública é exclusivamente fictícia; produção exige ambiente segregado e homologado."),
        ("Como evita fila?", "Capacidade por faixa, visão de ocupação e proteção concorrente no banco."),
    ], [3000, 6360])
    path = OUT / "Roteiro-Apresentacao-Produto-Saude-Perto-de-Voce.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(manual())
    print(pitch())
    print(compliance_report())
    print(presentation_script())
